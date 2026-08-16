import docx

doc = docx.Document('project_synopsis.docx')

word_count = 0
for p in doc.paragraphs:
    word_count += len(p.text.split())

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            word_count += len(cell.text.split())

print(f"Total Words in DOCX: {word_count}")
