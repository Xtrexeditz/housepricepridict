import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (margins) in dxa (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_runs(p, text):
    """Parses simple inline markdown styles (**bold**, *italic*, `code`) and adds to paragraph."""
    # Split text by bold markers (**), keeping the markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold segment
            bold_text = part[2:-2]
            # Check for nested italic
            sub_parts = re.split(r'(\*.*?\*)', bold_text)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    run = p.add_run(sub_part[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    run = p.add_run(sub_part)
                    run.bold = True
        else:
            # Regular segment, check for italic or inline code
            sub_parts = re.split(r'(\*.*?\*|`.*?`)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    run = p.add_run(sub_part[1:-1])
                    run.italic = True
                elif sub_part.startswith('`') and sub_part.endswith('`'):
                    run = p.add_run(sub_part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(99, 102, 241) # Indigo accent
                else:
                    p.add_run(sub_part)

def convert_md_to_docx(md_path, docx_path):
    print(f"Reading markdown from {md_path}...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    doc = docx.Document()
    
    # 1. Page Margins Setup (1 inch margins)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default paragraph format font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(31, 41, 55) # Dark gray body
    
    in_code_block = False
    code_text = []
    
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # --- CODE BLOCKS ---
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block, insert formatted container
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                # Apply light gray background & border simulation via left indent
                run = p.add_run("\n".join(code_text))
                run.font.name = 'Consolas'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(75, 85, 99)
                code_text = []
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_text.append(line.rstrip('\n'))
            i += 1
            continue
            
        # --- TABLES ---
        if stripped.startswith('|'):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            # End of table block, parse and generate docx table
            in_table = False
            
            # Filter divider lines (e.g. |:---|:---|)
            parsed_rows = []
            for t_line in table_lines:
                # Split cells and filter empty ends
                cells = [c.strip() for c in t_line.split('|')]
                if len(cells) > 1:
                    # Remove first and last empty cells due to leading/trailing |
                    if cells[0] == '': cells.pop(0)
                    if cells and cells[-1] == '': cells.pop()
                    
                    # Verify if it's a separator line
                    if cells and all(re.match(r'^:?-+:?$', c) for c in cells):
                        continue
                    parsed_rows.append(cells)
                    
            if parsed_rows:
                num_cols = max(len(row) for row in parsed_rows)
                t = doc.add_table(rows=len(parsed_rows), cols=num_cols)
                t.autofit = True
                
                # Style table
                for r_idx, row_data in enumerate(parsed_rows):
                    row = t.rows[r_idx]
                    is_header = (r_idx == 0)
                    
                    # Set heights
                    trPr = row._tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), '360' if is_header else '280')
                    trHeight.set(qn('w:hRule'), 'atLeast')
                    trPr.append(trHeight)
                    
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx >= num_cols: break
                        cell = row.cells[c_idx]
                        cell.text = "" # Clear default
                        
                        # Background coloring
                        if is_header:
                            set_cell_background(cell, "F3F4F6") # Light gray header
                        else:
                            if r_idx % 2 == 0:
                                set_cell_background(cell, "FAFAFA") # Alternating zebra
                                
                        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(4)
                        
                        add_paragraph_runs(p, cell_value)
                        
                        # Make header text bold and center it
                        if is_header:
                            for run in p.runs:
                                run.bold = True
                                
                doc.add_paragraph() # Add spacing after table
            table_lines = []
            
        # --- HEADINGS & PARAGRAPHS ---
        if stripped == "":
            i += 1
            continue
            
        # Title (Heading 1 main title)
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(17, 24, 39)
            
        # Subtitle / Section Heading 2
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(31, 41, 55)
            
        # Sub-section Heading 3
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:])
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(55, 65, 81)
            
        # Horizontal rule separator
        elif stripped == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("❖   ❖   ❖")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(156, 163, 175)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
        # Bullet list items
        elif stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_paragraph_runs(p, stripped[2:])
            
        # Standard paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_paragraph_runs(p, stripped)
            
        i += 1
        
    print(f"Saving compiled Word document to {docx_path}...")
    doc.save(docx_path)
    print("Word document generated successfully!")

if __name__ == '__main__':
    md = 'project_synopsis.md'
    docx_file = 'project_synopsis.docx'
    convert_md_to_docx(md, docx_file)
